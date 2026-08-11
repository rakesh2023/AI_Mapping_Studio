"""
AI Mapping Studio - local backend
==================================
A small Flask service that:
  1. Serves the existing static site (same origin -> no CORS needed).
  2. Exposes a REAL SQL Server connection API so the Metadata Explorer can
     list actual tables/columns from a live database via pyodbc.

Run:
    cd server
    pip install -r requirements.txt
    python app.py
Then open http://localhost:8000/

Security note: connection details are sent from the browser per request and
used only to open a short-lived connection; nothing is persisted server-side.
This is a local prototype helper, not a production gateway.
"""
import os
import json
from flask import Flask, request, jsonify, send_from_directory

try:
    import pyodbc
except ImportError:  # surfaced nicely to the UI instead of crashing on import
    pyodbc = None

try:
    import anthropic
except ImportError:  # AI generation degrades gracefully if the SDK isn't installed
    anthropic = None

# Optional file parsers for File System source systems. Each is guarded so the
# server still boots if one isn't installed; the extract endpoint reports which
# package is missing when a matching file type is uploaded.
try:
    import openpyxl
except ImportError:
    openpyxl = None

try:
    from pypdf import PdfReader
except ImportError:
    PdfReader = None

try:
    import docx  # python-docx
except ImportError:
    docx = None

# Static site lives one level up from /server
ROOT = os.path.abspath(os.path.join(os.path.dirname(__file__), ".."))

app = Flask(__name__, static_folder=None)


# --------------------------------------------------------------------------- #
#  Static file serving (index.html, pages/, css/, js/, data/, assets/)
# --------------------------------------------------------------------------- #
@app.route("/")
def index():
    return send_from_directory(ROOT, "index.html")


@app.route("/<path:path>")
def static_proxy(path):
    full = os.path.join(ROOT, path)
    if os.path.isfile(full):
        directory = os.path.dirname(full)
        return send_from_directory(directory, os.path.basename(full))
    return ("Not found", 404)


# --------------------------------------------------------------------------- #
#  Database API
# --------------------------------------------------------------------------- #
def build_connection_string(cfg):
    """Build a pyodbc connection string from the posted config."""
    driver = cfg.get("driver") or "ODBC Driver 17 for SQL Server"
    server = cfg.get("server", "")
    database = cfg.get("database", "")
    parts = [
        f"DRIVER={{{driver}}}",
        f"SERVER={server}",
        f"DATABASE={database}",
    ]
    if cfg.get("trusted"):  # Windows integrated auth
        parts.append("Trusted_Connection=yes")
    else:
        parts.append(f"UID={cfg.get('username','')}")
        parts.append(f"PWD={cfg.get('password','')}")
    # TrustServerCertificate is only understood by the modern "ODBC Driver xx
    # for SQL Server". The legacy "SQL Server" driver rejects it with
    # "Invalid connection string attribute", so only add it for MSODBC.
    if "odbc driver" in driver.lower():
        parts.append("TrustServerCertificate=yes")
    return ";".join(parts) + ";"


def open_connection(cfg):
    if pyodbc is None:
        raise RuntimeError("pyodbc is not installed on the server. Run: pip install -r requirements.txt")
    conn_str = build_connection_string(cfg)
    return pyodbc.connect(conn_str, timeout=int(cfg.get("timeout", 8)))


@app.route("/api/db/drivers")
def list_drivers():
    """Report which ODBC drivers are available on this machine."""
    if pyodbc is None:
        return jsonify(ok=False, error="pyodbc not installed", drivers=[])
    return jsonify(ok=True, drivers=[d for d in pyodbc.drivers()])


@app.route("/api/db/test", methods=["POST"])
def test_connection():
    cfg = request.get_json(force=True) or {}
    try:
        conn = open_connection(cfg)
        cur = conn.cursor()
        cur.execute("SELECT @@VERSION")
        version = cur.fetchone()[0]
        conn.close()
        return jsonify(ok=True, message="Connection successful.", version=version.split("\n")[0])
    except Exception as exc:  # noqa: BLE001 - surface any driver/auth error to UI
        return jsonify(ok=False, error=str(exc)), 400


@app.route("/api/db/metadata", methods=["POST"])
def get_metadata():
    """Return real tables + columns (with PK/FK) in the app's source-metadata shape."""
    cfg = request.get_json(force=True) or {}
    schema_filter = cfg.get("schema")  # optional, e.g. 'dbo'
    try:
        conn = open_connection(cfg)
        cur = conn.cursor()

        # Columns
        cur.execute(
            """
            SELECT c.TABLE_SCHEMA, c.TABLE_NAME, c.COLUMN_NAME, c.DATA_TYPE,
                   c.CHARACTER_MAXIMUM_LENGTH, c.NUMERIC_PRECISION,
                   c.IS_NULLABLE, c.COLUMN_DEFAULT, c.ORDINAL_POSITION
            FROM INFORMATION_SCHEMA.COLUMNS c
            JOIN INFORMATION_SCHEMA.TABLES t
              ON t.TABLE_SCHEMA = c.TABLE_SCHEMA AND t.TABLE_NAME = c.TABLE_NAME
            WHERE t.TABLE_TYPE = 'BASE TABLE'
            ORDER BY c.TABLE_SCHEMA, c.TABLE_NAME, c.ORDINAL_POSITION
            """
        )
        col_rows = cur.fetchall()

        # Primary keys
        cur.execute(
            """
            SELECT ku.TABLE_SCHEMA, ku.TABLE_NAME, ku.COLUMN_NAME
            FROM INFORMATION_SCHEMA.TABLE_CONSTRAINTS tc
            JOIN INFORMATION_SCHEMA.KEY_COLUMN_USAGE ku
              ON tc.CONSTRAINT_NAME = ku.CONSTRAINT_NAME
            WHERE tc.CONSTRAINT_TYPE = 'PRIMARY KEY'
            """
        )
        pks = {(r[0], r[1], r[2]) for r in cur.fetchall()}

        # Foreign keys
        cur.execute(
            """
            SELECT cu.TABLE_SCHEMA, cu.TABLE_NAME, cu.COLUMN_NAME
            FROM INFORMATION_SCHEMA.REFERENTIAL_CONSTRAINTS rc
            JOIN INFORMATION_SCHEMA.KEY_COLUMN_USAGE cu
              ON rc.CONSTRAINT_NAME = cu.CONSTRAINT_NAME
            """
        )
        fks = {(r[0], r[1], r[2]) for r in cur.fetchall()}

        # Row counts (best-effort, from sys catalog)
        rowcounts = {}
        try:
            cur.execute(
                """
                SELECT s.name, t.name, SUM(p.rows)
                FROM sys.tables t
                JOIN sys.schemas s ON s.schema_id = t.schema_id
                JOIN sys.partitions p ON p.object_id = t.object_id AND p.index_id IN (0,1)
                GROUP BY s.name, t.name
                """
            )
            for r in cur.fetchall():
                rowcounts[(r[0], r[1])] = int(r[2] or 0)
        except Exception:  # noqa: BLE001
            pass

        conn.close()

        tables = {}
        first_schema = None
        for r in col_rows:
            (tschema, tname, cname, dtype, charlen, numprec,
             is_nullable, default, _pos) = r
            if schema_filter and tschema != schema_filter:
                continue
            if first_schema is None:
                first_schema = tschema
            key = (tschema, tname)
            if key not in tables:
                tables[key] = {
                    "name": tname,
                    "schema": tschema,
                    "description": f"{tschema}.{tname}",
                    "rowCount": rowcounts.get(key, 0),
                    "columns": [],
                }
            tables[key]["columns"].append({
                "name": cname,
                "dataType": (dtype or "").upper(),
                "length": charlen if charlen not in (None, -1) else numprec,
                "nullable": (is_nullable == "YES"),
                "pk": (tschema, tname, cname) in pks,
                "fk": (tschema, tname, cname) in fks,
                "default": default,
                "description": "",
                "businessTerm": "",
                "sample": None,
                "distinctCount": None,
                "nullPct": 0,
            })

        table_list = list(tables.values())
        return jsonify(
            ok=True,
            connection=cfg.get("database", "Database"),
            schema=schema_filter or first_schema or "dbo",
            tableCount=len(table_list),
            columnCount=sum(len(t["columns"]) for t in table_list),
            tables=table_list,
        )
    except Exception as exc:  # noqa: BLE001
        return jsonify(ok=False, error=str(exc)), 400


def _quote(ident):
    """Safely quote a SQL Server identifier."""
    return "[" + str(ident).replace("]", "]]") + "]"


@app.route("/api/db/profile", methods=["POST"])
def profile_table():
    """Profile ONE table live: row count, and per-column null %, distinct
    count, min/max and top values. Body: connection cfg + {schema, table}."""
    cfg = request.get_json(force=True) or {}
    schema = cfg.get("schema") or "dbo"
    table = cfg.get("table")
    top_n = int(cfg.get("topN", 5))
    if not table:
        return jsonify(ok=False, error="No table specified."), 400
    try:
        conn = open_connection(cfg)
        cur = conn.cursor()
        fq = f"{_quote(schema)}.{_quote(table)}"

        # total rows
        cur.execute(f"SELECT COUNT(*) FROM {fq}")
        row_count = int(cur.fetchone()[0])

        # column list + types
        cur.execute(
            """
            SELECT COLUMN_NAME, DATA_TYPE, CHARACTER_MAXIMUM_LENGTH
            FROM INFORMATION_SCHEMA.COLUMNS
            WHERE TABLE_SCHEMA = ? AND TABLE_NAME = ?
            ORDER BY ORDINAL_POSITION
            """,
            schema, table,
        )
        col_defs = cur.fetchall()

        columns = []
        for cname, dtype, charlen in col_defs:
            col = _quote(cname)
            stats = {
                "name": cname,
                "dataType": (dtype or "").upper(),
                "length": charlen,
                "rowCount": row_count,
                "nullCount": 0,
                "nullPct": 0,
                "distinctCount": 0,
                "sample": None,
                "min": None,
                "max": None,
                "topValues": [],
            }
            try:
                cur.execute(
                    f"SELECT COUNT(*) - COUNT({col}), COUNT(DISTINCT {col}) FROM {fq}"
                )
                nulls, distinct = cur.fetchone()
                stats["nullCount"] = int(nulls or 0)
                stats["distinctCount"] = int(distinct or 0)
                stats["nullPct"] = round((int(nulls or 0) / row_count) * 100, 1) if row_count else 0

                # a sample non-null value
                cur.execute(f"SELECT TOP 1 {col} FROM {fq} WHERE {col} IS NOT NULL")
                s = cur.fetchone()
                if s and s[0] is not None:
                    stats["sample"] = str(s[0])[:120]

                # min / max for comparable types
                if (dtype or "").lower() not in ("text", "ntext", "image", "xml"):
                    try:
                        cur.execute(f"SELECT MIN({col}), MAX({col}) FROM {fq}")
                        mn, mx = cur.fetchone()
                        stats["min"] = None if mn is None else str(mn)[:60]
                        stats["max"] = None if mx is None else str(mx)[:60]
                    except Exception:  # noqa: BLE001
                        pass

                # top values only when the column is low-cardinality (looks categorical)
                if 0 < stats["distinctCount"] <= 50:
                    cur.execute(
                        f"SELECT TOP {top_n} {col} AS v, COUNT(*) AS c FROM {fq} "
                        f"WHERE {col} IS NOT NULL GROUP BY {col} ORDER BY c DESC"
                    )
                    for v, c in cur.fetchall():
                        pct = round((int(c) / row_count) * 100, 1) if row_count else 0
                        stats["topValues"].append({"value": str(v)[:60], "count": int(c), "pct": pct})
            except Exception:  # noqa: BLE001 - keep profiling other columns even if one fails
                pass
            columns.append(stats)

        conn.close()
        return jsonify(ok=True, schema=schema, table=table, rowCount=row_count, columns=columns)
    except Exception as exc:  # noqa: BLE001
        return jsonify(ok=False, error=str(exc)), 400


# --------------------------------------------------------------------------- #
#  AI mapping generation (real LLM via the Anthropic API)
# --------------------------------------------------------------------------- #
MAPPING_ITEM_SCHEMA = {
    "type": "object",
    "properties": {
        "mappings": {
            "type": "array",
            "items": {
                "type": "object",
                "properties": {
                    "targetEntity": {"type": "string"},
                    "targetColumn": {"type": "string"},
                    "sourceTable": {"type": "string"},
                    "sourceColumn": {"type": "string"},
                    "mappingType": {"type": "string", "enum": [
                        "Direct", "Derived", "Lookup", "Conditional", "Constant",
                        "Default", "Concatenation", "Split", "Format Conversion",
                        "Data Type Conversion", "Calculation", "Aggregation",
                        "Reference", "Custom", "Not Mapped"]},
                    "transformationRule": {"type": "string"},
                    "businessRule": {"type": "string"},
                    "nullHandling": {"type": "string"},
                    "confidence": {"type": "integer"},
                    "explanation": {"type": "string"},
                },
                "required": ["targetEntity", "targetColumn", "sourceTable", "sourceColumn",
                             "mappingType", "transformationRule", "businessRule",
                             "nullHandling", "confidence", "explanation"],
                "additionalProperties": False,
            },
        },
        "joins": {
            "type": "array",
            "items": {
                "type": "object",
                "properties": {
                    "targetEntity": {"type": "string"},
                    "joinCondition": {"type": "string"},
                },
                "required": ["targetEntity", "joinCondition"],
                "additionalProperties": False,
            },
        },
    },
    "required": ["mappings", "joins"],
    "additionalProperties": False,
}


def _ai_model():
    """Model id for the active endpoint. Corporate/Bedrock gateways expect their
    own ids (e.g. bedrock.anthropic.claude-opus-4-8). Read from env, but strip any
    context-window suffix like '[1m]' — this key rejects the suffixed variant."""
    m = (os.environ.get("AIMS_MODEL")
         or os.environ.get("ANTHROPIC_DEFAULT_OPUS_MODEL")
         or "claude-opus-5")
    return m.split("[", 1)[0]


def _ca_bundle():
    """Locate a CA bundle that trusts the corporate TLS-intercepting proxy.
    Prefer the Windows-trust-store bundle generated next to this file (built by
    the setup step), then fall back to the env-configured bundles."""
    here = os.path.dirname(__file__)
    candidates = [
        os.path.join(here, "win-ca-bundle.pem"),
        os.environ.get("SSL_CERT_FILE"),
        os.environ.get("REQUESTS_CA_BUNDLE"),
        os.environ.get("CURL_CA_BUNDLE"),
    ]
    for c in candidates:
        if c and os.path.isfile(c):
            return c
    return None


def _anthropic_client():
    """Build a client whose httpx transport trusts the corporate CA bundle, so the
    TLS handshake to the internal gateway (ANTHROPIC_BASE_URL) succeeds instead of
    failing with 'Connection error'. base_url + auth token are read from env."""
    kwargs = {}
    ca = _ca_bundle()
    if ca:
        try:
            import httpx
            kwargs["http_client"] = httpx.Client(verify=ca, timeout=600.0)
        except Exception:  # noqa: BLE001 - fall back to the default client
            pass
    return anthropic.Anthropic(**kwargs)


@app.route("/api/ai/status")
def ai_status():
    """Report whether the AI backend is usable (SDK present + credentials resolvable)."""
    if anthropic is None:
        return jsonify(ok=False, reason="The 'anthropic' SDK is not installed (pip install anthropic).")
    if not (os.environ.get("ANTHROPIC_API_KEY") or os.environ.get("ANTHROPIC_AUTH_TOKEN")):
        return jsonify(ok=False, reason="No ANTHROPIC_API_KEY / ANTHROPIC_AUTH_TOKEN set in the server environment.")
    base = os.environ.get("ANTHROPIC_BASE_URL")
    return jsonify(ok=True, model=_ai_model(), endpoint=(base or "api.anthropic.com"))


@app.route("/api/ai/generate-mappings", methods=["POST"])
def generate_mappings():
    """Use Claude to map uploaded target-schema fields to live source columns.

    Body: {
      source: {connection, schema, tables:[{name, columns:[{name,dataType,...}]}]},
      targetEntities: [{name, table, fields:[{name,dataType,length,mandatory,pk,fk,
                        fkReference,accepted,description}]}],
      businessContext: str, instructions: str, strategy: str
    }
    """
    if anthropic is None:
        return jsonify(ok=False, error="The 'anthropic' SDK is not installed on the server. Run: pip install anthropic"), 400
    body = request.get_json(force=True) or {}
    source = body.get("source") or {}
    target_entities = body.get("targetEntities") or []
    if not source.get("tables"):
        return jsonify(ok=False, error="No source tables provided. Load a source system first."), 400
    if not target_entities:
        return jsonify(ok=False, error="No target entities provided. Upload a target schema first."), 400

    # Group the source columns BY TABLE with numbered headers so the model can see and
    # scan every table, instead of anchoring on the first entries of a flat list.
    src_tables = source["tables"]
    blocks = []
    for i, t in enumerate(src_tables, 1):
        col_lines = []
        for c in t.get("columns", []):
            extra_bits = []
            if c.get("businessTerm"): extra_bits.append("business term: " + str(c["businessTerm"]))
            if c.get("description"): extra_bits.append(str(c["description"]))
            if c.get("sample") not in (None, ""): extra_bits.append("e.g. " + str(c["sample"]))
            col_lines.append(f"    {c['name']} ({c.get('dataType','')}"
                             + (f"({c['length']})" if c.get('length') else "") + ")"
                             + (" — " + "; ".join(extra_bits) if extra_bits else ""))
        blocks.append(f"[{i}/{len(src_tables)}] TABLE {t['name']} ({len(t.get('columns', []))} columns):\n"
                      + "\n".join(col_lines))
    source_block = "\n\n".join(blocks)
    src_table_names = ", ".join(t["name"] for t in src_tables)

    def _target_block(entities):
        tgt_blocks = []
        for e in entities:
            rows = []
            for f in e.get("fields", []):
                attrs = []
                if f.get("mandatory"): attrs.append("mandatory")
                if f.get("pk"): attrs.append("PK")
                if f.get("fk"): attrs.append("FK" + (f" -> {f['fkReference']}" if f.get("fkReference") else ""))
                if f.get("accepted"): attrs.append("accepted: " + str(f["accepted"]))
                rows.append(f"  - {f['name']} ({f.get('dataType','')}"
                            + (f"({f['length']})" if f.get('length') else "") + ")"
                            + (" [" + ", ".join(attrs) + "]" if attrs else "")
                            + (f" — {f['description']}" if f.get("description") else ""))
            tgt_blocks.append(f"Entity {e['name']} (table {e.get('table','')}):\n" + "\n".join(rows))
        return "\n\n".join(tgt_blocks)

    strategy = body.get("strategy", "Balanced")
    biz = (body.get("businessContext") or "").strip()
    extra = (body.get("instructions") or "").strip()

    system = (
        "You are a senior data-migration mapping engineer. You produce precise "
        "source-to-target field mappings for a database migration. For every target "
        "field you are given, choose the single best source column (from the provided "
        "source column list only — never invent a source column). Decide the mapping "
        "type, write a concrete transformation rule (SQL-like), a short business rule, "
        "null handling, and a 0-100 confidence score.\n\n"
        "MATCHING RULES — source and target names WILL differ; match on meaning, not "
        "exact strings:\n"
        "- SEARCH ACROSS ALL SOURCE TABLES. The source columns are grouped under "
        "numbered TABLE headers ([1/N] ... [N/N]); the best match for a target field "
        "is frequently in a table whose NAME looks unrelated to the target entity. "
        "Do NOT restrict yourself to the first table or to a table whose name resembles "
        "the target — scan every table's columns before deciding, and before marking a "
        "field 'Not Mapped' confirm no column in ANY listed table fits.\n"
        "- Normalize names before comparing: ignore case, and treat snake_case, "
        "camelCase, PascalCase and kebab-case as equivalent (POLICY_NUMBER == "
        "PolicyNumber == policyNumber).\n"
        "- Expand and normalize common abbreviations both ways: CUST/CUSTOMER, "
        "NBR/NUM/NO/# = number, DT/DATE, AMT = amount, ADDR = address, DESC = "
        "description, CD/CODE, ID/IDENTIFIER, FNAME/FIRST_NAME, LNAME/LAST_NAME, "
        "DOB = date of birth, TS = timestamp, QTY = quantity, PCT = percent, "
        "STS/STATUS, TEL/PH = phone, EMAIL/EMAIL_ADDR, ZIP/POSTAL_CODE, CTRY/COUNTRY, "
        "ST = state, ORG = organization, ACCT = account, AGT = agent, TXN = transaction.\n"
        "- Use each column's business term, description and sample value (when given) "
        "as strong matching signals — a matching business term outweighs a differing "
        "column name.\n"
        "- Consider data-type compatibility (a date target should map from a date/"
        "datetime/timestamp source; a numeric amount from a numeric or numeric-text "
        "column via conversion).\n"
        "- Prefer a same-named column in a differently-named table over a poor name "
        "match in another table; the table names need not align.\n"
        "- Pick the best candidate even when the name overlap is partial; lower the "
        "confidence score to reflect uncertainty rather than refusing to map.\n\n"
        "If, after applying all rules above, no plausible source column exists, use "
        "mappingType 'Not Mapped', set sourceTable and sourceColumn to empty strings, "
        "confidence 0, and explain the gap. Do NOT mark a field 'Not Mapped' merely "
        "because the names are spelled differently.\n\n"
        f"Apply the '{strategy}' strategy: Conservative = only map high-confidence "
        "matches; Balanced = map likely matches and flag uncertain ones; Aggressive = "
        "map as many as possible including low-confidence guesses."
    )
    system += (
        "\n\nJOIN CONDITIONS: A target entity is often populated by combining several "
        "source tables. For EACH target entity, determine the SQL JOIN that assembles "
        "the source rows feeding its fields. Infer join keys from primary/foreign keys, "
        "matching *_ID / *_CD / *_NBR columns, and shared business terms across the "
        "tables you actually used in that entity's mappings. Write a runnable SQL "
        "snippet, e.g. 'FROM CLM_TXN c JOIN PARTY_MST p ON c.PARTY_ID = p.PARTY_ID'. If "
        "the entity draws from a single table, give just its FROM clause "
        "('FROM CLM_TXN'). If no source tables were used, return an empty string."
    )
    # Ask for JSON in the prompt too, so we don't depend on structured-output
    # support (internal/Bedrock gateways may not accept output_config.format).
    system += (" Respond with ONLY a JSON object of the form "
               '{"mappings": [ ... ], "joins": [ ... ]}. Each mappings item has keys '
               "targetEntity, targetColumn, sourceTable, sourceColumn, mappingType, "
               "transformationRule, businessRule, nullHandling, confidence "
               "(integer 0-100), explanation. Each joins item has keys targetEntity and "
               "joinCondition (the SQL FROM/JOIN snippet described above), one per target "
               "entity. No prose, no markdown fences.")

    model = _ai_model()

    def _build_user(entities):
        u = ("SOURCE DATABASE: " + str(source.get("connection", "")) + " — "
             + str(len(src_tables)) + " tables to search: " + src_table_names + "\n\n"
             + "SOURCE COLUMNS (grouped by table; search ALL of them):\n" + source_block
             + "\n\nTARGET FIELDS TO MAP:\n" + _target_block(entities))
        if biz:
            u += "\n\nBUSINESS CONTEXT:\n" + biz
        if extra:
            u += "\n\nADDITIONAL INSTRUCTIONS:\n" + extra
        u += ("\n\nReturn one mapping object per target field listed above. For each "
              "field, scan every one of the " + str(len(src_tables)) + " source tables "
              "before choosing the best column or marking it Not Mapped.")
        return u

    def _call_model(entities):
        """One model call for the given target entities. Returns the parsed dict."""
        client = _anthropic_client()
        base_kwargs = dict(model=model, max_tokens=16000, system=system,
                           messages=[{"role": "user", "content": _build_user(entities)}])

        def run(extra_cfg):
            with client.messages.stream(**base_kwargs, **extra_cfg) as stream:
                return stream.get_final_message()

        attempts = [
            {"output_config": {"effort": "medium", "format": {"type": "json_schema", "schema": MAPPING_ITEM_SCHEMA}}},
            {"output_config": {"format": {"type": "json_schema", "schema": MAPPING_ITEM_SCHEMA}}},
            {"output_config": {"effort": "medium"}},
            {},
        ]
        resp, last_err = None, None
        for cfg in attempts:
            try:
                resp = run(cfg); break
            except Exception as e:  # noqa: BLE001
                last_err = e
        if resp is None:
            raise last_err
        if getattr(resp, "stop_reason", None) == "refusal":
            raise RuntimeError("The request was declined by safety classifiers.")
        text = next((b.text for b in resp.content if getattr(b, "type", None) == "text"), "")
        usage = getattr(resp, "usage", None)
        return _parse_mapping_json(text), usage

    # A single target table can have many columns (100+). One mapping object per field
    # is large, so a big table's JSON response can exceed the output token limit and get
    # truncated -> parse fails -> 0 mappings for that table. Split a table's fields into
    # chunks and call the model per chunk, then aggregate.
    FIELD_CHUNK = 40

    def _chunk_entity(e):
        fields = e.get("fields", []) or []
        if len(fields) <= FIELD_CHUNK:
            return [e]
        chunks = []
        for start in range(0, len(fields), FIELD_CHUNK):
            sub = dict(e)
            sub["fields"] = fields[start:start + FIELD_CHUNK]
            chunks.append(sub)
        return chunks

    try:
        # Process ONE target entity (in field-chunks) per model call. This keeps each
        # request small and fast so neither many tables nor a single wide table overflows
        # the output token limit or times out the gateway.
        by_key, joins_in = {}, {}
        in_tokens = out_tokens = 0
        for e in target_entities:
            for chunk in _chunk_entity(e):
                data, usage = _call_model([chunk])
                for m in (data.get("mappings", []) or []):
                    k = ((m.get("targetEntity") or "").strip(), (m.get("targetColumn") or "").strip())
                    if k not in by_key:
                        by_key[k] = m
                for j in (data.get("joins") or []):
                    if isinstance(j, dict) and (j.get("targetEntity") or "").strip():
                        joins_in.setdefault((j["targetEntity"]).strip(), (j.get("joinCondition") or "").strip())
                if usage:
                    in_tokens += getattr(usage, "input_tokens", 0) or 0
                    out_tokens += getattr(usage, "output_tokens", 0) or 0

        # Guarantee one row per requested target field, so every selected table always
        # shows up in the workspace — even when the AI found no matches at all.
        returned_count = 0
        mappings = []
        for e in target_entities:
            ename = e.get("name", "")
            for f in e.get("fields", []):
                cname = f.get("name", "")
                m = by_key.get((ename, cname))
                if m:
                    m.setdefault("targetEntity", ename)
                    m.setdefault("targetColumn", cname)
                    mappings.append(m)
                    returned_count += 1
                else:
                    mappings.append({
                        "targetEntity": ename, "targetColumn": cname,
                        "sourceTable": "", "sourceColumn": "",
                        "mappingType": "Not Mapped", "transformationRule": "",
                        "businessRule": "No matching source column was found for this field.",
                        "nullHandling": "N/A", "confidence": 0,
                        "explanation": "The AI did not return a mapping for this target field; flagged for manual review.",
                    })

        joins = [{"targetEntity": e.get("name", ""),
                  "joinCondition": joins_in.get((e.get("name", "") or "").strip(), "")}
                 for e in target_entities]

        return jsonify(ok=True, model=model, mappings=mappings, joins=joins,
                       returnedCount=returned_count,
                       usage={"input_tokens": in_tokens, "output_tokens": out_tokens})
    except Exception as exc:  # noqa: BLE001
        import traceback
        traceback.print_exc()
        msg = str(exc) or (exc.__class__.__name__ + " (see server log)")
        return jsonify(ok=False, error=msg), 400


SINGLE_MAPPING_SCHEMA = {
    "type": "object",
    "properties": {
        "sourceTable": {"type": "string"},
        "sourceColumn": {"type": "string"},
        "mappingType": {"type": "string", "enum": [
            "Direct", "Derived", "Lookup", "Conditional", "Constant", "Default",
            "Concatenation", "Split", "Format Conversion", "Data Type Conversion",
            "Calculation", "Aggregation", "Reference", "Custom", "Not Mapped"]},
        "transformationRule": {"type": "string"},
        "businessRule": {"type": "string"},
        "lookupTable": {"type": "string"},
        "defaultValue": {"type": "string"},
        "nullHandling": {"type": "string"},
        "confidence": {"type": "integer"},
        "explanation": {"type": "string"},
        "joinCondition": {"type": "string"},
    },
    "required": ["sourceTable", "sourceColumn", "mappingType", "transformationRule",
                 "businessRule", "nullHandling", "confidence", "explanation"],
    "additionalProperties": False,
}


@app.route("/api/ai/regenerate-mapping", methods=["POST"])
def regenerate_mapping():
    """Re-map a SINGLE target field with Claude, honoring the user's instruction
    (e.g. 'hardcode currency as USD'). Body: {mapping:{...current row...},
    sourceColumns:[{table,column,dataType}], instruction:str}."""
    if anthropic is None:
        return jsonify(ok=False, error="The 'anthropic' SDK is not installed on the server."), 400
    body = request.get_json(force=True) or {}
    m = body.get("mapping") or {}
    src_cols = body.get("sourceColumns") or []
    instruction = (body.get("instruction") or "").strip()
    current_join = (body.get("currentJoin") or "").strip()
    entity_tables = body.get("entitySourceTables") or []   # source tables used by this entity
    if not m.get("targetColumn"):
        return jsonify(ok=False, error="No target field provided."), 400

    src_block = "\n".join(f"{c.get('table','')}.{c.get('column','')} ({c.get('dataType','')})"
                          for c in src_cols) or "(no source columns provided)"
    system = (
        "You are a senior data-migration mapping engineer. Re-map ONE target field, "
        "following the user's instruction exactly. The instruction takes priority over "
        "your default choice — if the user says to hardcode a constant, set mappingType "
        "to 'Constant', put the value in defaultValue, and write transformationRule like "
        "CONSTANT('<value>'). If they specify a lookup, use 'Lookup' and fill lookupTable. "
        "Choose sourceTable/sourceColumn from the provided source list only (leave both "
        "empty for Constant/Default/Not Mapped).\n\n"
        "JOIN CONDITION: You are also given the target entity's CURRENT join (FROM/JOIN "
        "SQL that assembles its source rows) and the source tables it already uses. If your "
        "chosen sourceTable is NOT already in that FROM/JOIN, UPDATE joinCondition to add a "
        "JOIN for it — infer the join key from matching *_ID / *_CD / *_NBR columns or shared "
        "business keys between the tables. If the chosen table is already covered (or the "
        "mapping is Constant/Default/Not Mapped), return the current join unchanged. Write a "
        "runnable snippet, e.g. 'FROM CLM_TXN c JOIN PARTY_MST p ON c.PARTY_ID = p.PARTY_ID'.\n\n"
        "Return the full updated mapping. Respond with ONLY a JSON object with keys "
        "sourceTable, sourceColumn, mappingType, transformationRule, businessRule, "
        "lookupTable, defaultValue, nullHandling, confidence (0-100 integer), explanation, "
        "joinCondition. No prose, no markdown fences."
    )
    user = (
        "TARGET FIELD: " + str(m.get("targetEntity", "")) + "." + str(m.get("targetColumn", ""))
        + " (" + str(m.get("targetDataType", "")) + ")\n"
        + "CURRENT MAPPING: " + json.dumps({k: m.get(k) for k in
            ("sourceTable", "sourceColumn", "mappingType", "transformationRule",
             "businessRule", "lookupTable", "defaultValue", "nullHandling")}, ensure_ascii=False)
        + "\n\nENTITY '" + str(m.get("targetEntity", "")) + "' CURRENT JOIN: "
        + (current_join or "(none yet)")
        + "\nSOURCE TABLES ALREADY USED BY THIS ENTITY: "
        + (", ".join(entity_tables) if entity_tables else "(none)")
        + "\n\nAVAILABLE SOURCE COLUMNS:\n" + src_block
        + "\n\nUSER INSTRUCTION (apply this): " + (instruction or "Improve this mapping.")
    )

    model = _ai_model()
    try:
        client = _anthropic_client()
        base_kwargs = dict(model=model, max_tokens=2000, system=system,
                           messages=[{"role": "user", "content": user}])

        def run(extra):
            with client.messages.stream(**base_kwargs, **extra) as stream:
                return stream.get_final_message()
        try:
            resp = run({"output_config": {"format": {"type": "json_schema", "schema": SINGLE_MAPPING_SCHEMA}}})
        except Exception:  # noqa: BLE001
            resp = run({})

        if getattr(resp, "stop_reason", None) == "refusal":
            return jsonify(ok=False, error="The request was declined by safety classifiers."), 400
        text = next((b.text for b in resp.content if getattr(b, "type", None) == "text"), "")
        parsed = _parse_mapping_json(text)
        # _parse_mapping_json returns {"mappings":[]} on failure; a single object is what we want.
        if isinstance(parsed, dict) and "mappings" in parsed and not parsed.get("mappingType"):
            parsed = {}
        return jsonify(ok=True, model=model, mapping=parsed)
    except Exception as exc:  # noqa: BLE001
        import traceback
        traceback.print_exc()
        return jsonify(ok=False, error=(str(exc) or exc.__class__.__name__)), 400


def _parse_mapping_json(text):
    """Best-effort JSON extraction — tolerates ```json fences or leading prose."""
    text = (text or "").strip()
    try:
        return json.loads(text)
    except Exception:  # noqa: BLE001
        pass
    start = text.find("{")
    end = text.rfind("}")
    if start != -1 and end != -1 and end > start:
        try:
            return json.loads(text[start:end + 1])
        except Exception:  # noqa: BLE001
            pass
    return {"mappings": []}


# --------------------------------------------------------------------------- #
#  File System source extraction (Excel / PDF / Word / SQL -> source schema)
# --------------------------------------------------------------------------- #
EXTRACT_TEXT_BUDGET = 60000   # hard cap chars per chunk (avoids input truncation)
EXTRACT_AI_CHUNK = 6000       # target chars per AI text-chunk — small enough that the
                              # model returns EVERY table in the slice instead of
                              # summarising a long list into a few tables.
EXTRACT_XLSX_ROW_CAP = 500     # rows per sheet-slice (tall sheets)
EXTRACT_XLSX_COL_CAP = 150     # columns per sheet-slice (WIDE sheets)
EXTRACT_XLSX_SAMPLE_ROWS = 8   # sample data rows per column-slice (wide sheets)
EXTRACT_MAX_CHUNKS = 200       # safety cap on total model calls per file


def _xlsx_sheet_chunks(title, grid):
    """Turn one sheet's cell grid (list of row-lists) into text chunks.

    - Narrow sheet  -> row-slices (repeat header) so tall sheets are fully read.
    - WIDE sheet    -> column-slices (each carries its columns' header + a few sample
      rows), so a sheet with thousands of COLUMNS never overflows the input/output.
    Chunks for the same sheet all use the same 'Sheet:' name, so the merge step unions
    their columns back into one table.
    """
    if not grid:
        return []
    ncols = max((len(r) for r in grid), default=0)
    chunks = []

    if ncols > EXTRACT_XLSX_COL_CAP:
        # ---- WIDE sheet: slice by columns ----
        sample = grid[:EXTRACT_XLSX_SAMPLE_ROWS]   # header + a few rows for type inference
        for cstart in range(0, ncols, EXTRACT_XLSX_COL_CAP):
            cend = min(cstart + EXTRACT_XLSX_COL_CAP, ncols)
            lines = []
            for r in sample:
                cells = [(r[i] if i < len(r) else "") for i in range(cstart, cend)]
                if any(cells):
                    lines.append("\t".join(cells))
            if lines:
                chunks.append("Sheet: " + title + " (columns " + str(cstart + 1) + "-" +
                              str(cend) + " of " + str(ncols) + ")\n" + "\n".join(lines))
        return chunks

    # ---- Narrow sheet ----
    grid = [r for r in grid if any(r)]
    if not grid:
        return []
    header_cells = grid[0]
    header = "\t".join(header_cells)

    # Detect a data-dictionary layout: a column whose header names the TABLE/ENTITY.
    # If found, GROUP data rows by that table so each chunk holds complete tables
    # (raw fixed-size row blocks confuse the model and it returns almost nothing).
    tbl_col = None
    for i, h in enumerate(header_cells):
        hl = str(h).strip().lower().replace("_", "").replace(" ", "")
        if hl in ("table", "tablename", "entity", "entityname", "objectname", "object"):
            tbl_col = i; break

    if tbl_col is not None:
        # group rows by table value, preserving order
        groups, gorder = {}, []
        for r in grid[1:]:
            tv = (r[tbl_col] if tbl_col < len(r) else "").strip()
            if not tv:
                continue
            if tv not in groups:
                groups[tv] = []; gorder.append(tv)
            groups[tv].append("\t".join(r))
        if gorder:
            # batch a few tables per chunk, bounded by char budget
            cur, cur_len, count = [], 0, 0
            def flush():
                if cur:
                    chunks.append("Sheet: " + title +
                                  " (data dictionary — each row is a COLUMN; the '" +
                                  str(header_cells[tbl_col]) + "' cell names its table)\n" +
                                  header + "\n" + "\n".join(cur))
            for tv in gorder:
                block_rows = groups[tv]
                block_text = "\n".join(block_rows)
                if cur and (count >= 6 or cur_len + len(block_text) > EXTRACT_AI_CHUNK):
                    flush(); cur, cur_len, count = [], 0, 0
                cur.extend(block_rows); cur_len += len(block_text); count += 1
            flush()
            return chunks

    # No table column -> plain row-slices (small blocks, header repeated).
    rows = ["\t".join(r) for r in grid]
    block = 120
    if len(rows) <= block:
        return ["Sheet: " + title + "\n" + "\n".join(rows)]
    for start in range(1, len(rows), block):
        part = [header] + rows[start:start + block]
        chunks.append("Sheet: " + title + " (rows " + str(start) + "-" +
                      str(start + len(part) - 2) + ")\n" + "\n".join(part))
    return chunks


def _split_text_chunks(text, size=EXTRACT_TEXT_BUDGET):
    """Split a long text into <=size chunks, breaking on newlines where possible."""
    text = text or ""
    if len(text) <= size:
        return [text] if text.strip() else []
    chunks, i, n = [], 0, len(text)
    while i < n:
        end = min(i + size, n)
        if end < n:
            nl = text.rfind("\n", i, end)   # prefer a clean line break
            if nl > i + int(size * 0.5):
                end = nl
        chunk = text[i:end]
        if chunk.strip():
            chunks.append(chunk)
        i = end
    return chunks


# A new table usually starts with one of these markers in a dictionary / DDL / doc.
_TABLE_MARKER = None
def _table_marker():
    global _TABLE_MARKER
    if _TABLE_MARKER is None:
        import re
        _TABLE_MARKER = re.compile(
            r"^\s*(?:CREATE\s+TABLE\b|TABLE\s*[:\-]|TABLE\s+NAME\s*[:\-]|ENTITY\s*[:\-]|"
            r"\d+[.)]\s*TABLE\b)", re.IGNORECASE)
    return _TABLE_MARKER


def _split_by_tables(text, tables_per_chunk=8, max_chars=EXTRACT_AI_CHUNK):
    """Split dictionary/DDL text on TABLE boundaries and batch a small number of tables
    per chunk, so each AI call is asked for only a few tables (it then returns them all
    instead of summarising a long list). Falls back to size-splitting when no markers.
    """
    text = text or ""
    lines = text.split("\n")
    marker = _table_marker()
    # find line indexes where a new table begins
    starts = [i for i, ln in enumerate(lines) if marker.match(ln)]
    if len(starts) < 2:
        return _split_text_chunks(text, max_chars)   # not a table-per-block layout

    # build per-table blocks
    blocks = []
    for bi, s in enumerate(starts):
        e = starts[bi + 1] if bi + 1 < len(starts) else len(lines)
        # keep any preamble before the first table with the first block
        if bi == 0 and s > 0:
            s = 0
        blocks.append("\n".join(lines[s:e]).strip())
    blocks = [b for b in blocks if b]

    # group blocks into chunks by count AND size
    chunks, cur, cur_len = [], [], 0
    for b in blocks:
        if cur and (len(cur) >= tables_per_chunk or cur_len + len(b) > max_chars):
            chunks.append("\n\n".join(cur)); cur, cur_len = [], 0
        cur.append(b); cur_len += len(b) + 2
    if cur:
        chunks.append("\n\n".join(cur))
    return chunks


def _extract_file_chunks(filename, raw):
    """Extract text from an uploaded file as a LIST of chunks for looped extraction.

    Big files lose tables when sent to the model in one shot (input/output truncation
    and summarising). So we slice the file — Excel by sheet (further split if huge),
    PDF by page-groups, others by text size — and the caller runs the AI per chunk and
    merges the results. Returns (chunks, error); one is always None.
    """
    name = (filename or "").lower()
    try:
        if name.endswith((".xlsx", ".xlsm", ".xls")):
            if openpyxl is None:
                return None, "The 'openpyxl' package is not installed on the server (pip install openpyxl)."
            import io
            wb = openpyxl.load_workbook(io.BytesIO(raw), read_only=True, data_only=True)
            chunks = []
            for ws in wb.worksheets:
                # Read the sheet into a cell grid, then chunk by rows (tall) or columns
                # (wide) via _xlsx_sheet_chunks so no sheet overflows the model limits.
                grid = []
                for row in ws.iter_rows(values_only=True):
                    grid.append(["" if c is None else str(c) for c in row])
                chunks.extend(_xlsx_sheet_chunks(ws.title, grid))
            return chunks, None

        if name.endswith(".pdf"):
            if PdfReader is None:
                return None, "The 'pypdf' package is not installed on the server (pip install pypdf)."
            import io
            reader = PdfReader(io.BytesIO(raw))
            full = "\n".join((p.extract_text() or "") for p in reader.pages)
            # Split on TABLE boundaries (batch a few tables per chunk) so the model
            # returns EVERY table instead of summarising a long dictionary into a few.
            return _split_by_tables(full), None

        if name.endswith(".docx"):
            if docx is None:
                return None, "The 'python-docx' package is not installed on the server (pip install python-docx)."
            import io
            document = docx.Document(io.BytesIO(raw))
            parts = [p.text for p in document.paragraphs if p.text]
            for tbl in document.tables:
                for row in tbl.rows:
                    cells = [c.text for c in row.cells]
                    if any(cells):
                        parts.append("\t".join(cells))
            return _split_by_tables("\n".join(parts)), None

        # .sql, .txt, .csv, .json, .xml and anything else -> decode as text, then split by
        # table boundaries (SQL CREATE TABLE with real DDL is handled by the parser upstream).
        return _split_by_tables(raw.decode("utf-8", errors="ignore")), None
    except Exception as exc:  # noqa: BLE001
        return None, "Could not read the file: " + (str(exc) or exc.__class__.__name__)


def _extract_file_text(filename, raw):
    """Back-compat: return the whole file text joined (used by the SQL fast-path check)."""
    chunks, err = _extract_file_chunks(filename, raw)
    if err:
        return None, err
    return "\n".join(chunks or []), None


_SQL_TYPE_KEYWORDS = (
    "int", "bigint", "smallint", "tinyint", "bit", "decimal", "numeric", "money",
    "smallmoney", "float", "real", "date", "datetime", "datetime2", "smalldatetime",
    "time", "timestamp", "char", "varchar", "nchar", "nvarchar", "text", "ntext",
    "binary", "varbinary", "image", "uniqueidentifier", "xml", "number", "varchar2",
    "nvarchar2", "clob", "blob", "boolean", "bool", "json", "double", "serial",
)
# non-column lines inside a CREATE TABLE (...) body
_SQL_CONSTRAINT_STARTS = (
    "primary", "foreign", "unique", "constraint", "check", "key", "index",
    "period", ")", "with", "on",
)


def _parse_sql_ddl(text):
    """Deterministically parse CREATE TABLE statements into the source shape.

    LLMs tend to SUMMARISE large/repetitive DDL scripts (dropping tables), so for
    SQL scripts we parse every CREATE TABLE ourselves — no truncation, every table.
    Returns a list of {name, columns:[...]} (possibly empty if nothing parsed).
    """
    import re
    tables = []
    # Match: CREATE TABLE [schema.]name ( ... )  — capture the parenthesised body.
    pattern = re.compile(
        r"CREATE\s+TABLE\s+(?:IF\s+NOT\s+EXISTS\s+)?([^\s(]+)\s*\((.*?)\)\s*;",
        re.IGNORECASE | re.DOTALL,
    )
    for m in pattern.finditer(text):
        raw_name = m.group(1).strip().strip('`"[]')
        # keep only the final identifier (drop schema/db qualifiers)
        name = raw_name.split(".")[-1].strip('`"[]')
        body = m.group(2)

        # Split the body on top-level commas (ignore commas inside type parens like DECIMAL(18,2)).
        cols, depth, buf = [], 0, []
        for ch in body:
            if ch == "(":
                depth += 1; buf.append(ch)
            elif ch == ")":
                depth -= 1; buf.append(ch)
            elif ch == "," and depth == 0:
                cols.append("".join(buf)); buf = []
            else:
                buf.append(ch)
        if buf:
            cols.append("".join(buf))

        columns = []
        for raw_col in cols:
            line = raw_col.strip().strip(",").strip()
            if not line:
                continue
            low = line.lower()
            if low.startswith(_SQL_CONSTRAINT_STARTS):
                continue
            parts = line.split(None, 2)   # col_name, type, rest
            if len(parts) < 2:
                continue
            col_name = parts[0].strip('`"[]')
            type_token = parts[1]
            base_type = type_token.split("(")[0].lower()
            if base_type not in _SQL_TYPE_KEYWORDS:
                # not a recognisable column definition (e.g. a stray constraint)
                continue
            length = None
            lm = re.search(r"\(\s*(\d+)", type_token)
            if lm:
                length = int(lm.group(1))
            columns.append({
                "name": col_name, "dataType": base_type, "length": length,
                "businessTerm": "", "description": "", "sample": "",
            })
        if name and columns:
            tables.append({"name": name, "columns": columns})
    return tables


# Header synonyms for a structured Excel data dictionary (normalised: lowercase,
# no spaces/underscores). Used to read attributes DIRECTLY from cells — no AI.
_XLSX_HDR = {
    "table":       ("table", "tablename", "targettable", "physicaltable", "entity",
                    "entityname", "objectname", "object", "sourcetable"),
    "column":      ("column", "columnname", "field", "fieldname", "attribute",
                    "attributename", "sourcecolumn", "targetcolumn"),
    "datatype":    ("datatype", "type", "columntype", "sqltype", "fieldtype"),
    "length":      ("length", "len", "size", "columnlength", "fieldlength", "maxlength"),
    "description": ("description", "desc", "comment", "comments", "notes", "definition", "remarks"),
    "businessterm":("businessterm", "business", "glossaryterm", "businessname", "term"),
    "sample":      ("sample", "samplevalue", "example", "examplevalue", "sampledata"),
}


def _norm_hdr(h):
    return str(h or "").strip().lower().replace("_", "").replace(" ", "").replace("-", "")


def _parse_xlsx_dictionary(raw):
    """Deterministically parse a STRUCTURED Excel data dictionary into the source shape,
    reading every attribute (name, dataType, length, description, businessTerm, sample)
    straight from the cells — no AI, verbatim, instant.

    Returns a list of tables, or None if the sheet isn't a recognisable dictionary
    (no table-name + column-name header pair) so the caller falls back to the AI loop.
    """
    if openpyxl is None:
        return None
    import io
    try:
        wb = openpyxl.load_workbook(io.BytesIO(raw), read_only=True, data_only=True)
    except Exception:  # noqa: BLE001
        return None

    tables, order = {}, []
    recognised_any = False

    for ws in wb.worksheets:
        rows = [["" if c is None else str(c).strip() for c in r]
                for r in ws.iter_rows(values_only=True)]
        rows = [r for r in rows if any(v for v in r)]
        if len(rows) < 2:
            continue
        header = rows[0]
        # map each attribute to a column index via synonyms
        idx = {}
        for ci, h in enumerate(header):
            n = _norm_hdr(h)
            for key, syns in _XLSX_HDR.items():
                if key not in idx and n in syns:
                    idx[key] = ci
                    break
        # need at least a TABLE column and a COLUMN column to be a dictionary
        if "table" not in idx or "column" not in idx:
            continue
        recognised_any = True

        def cell(r, key):
            i = idx.get(key)
            return (r[i].strip() if (i is not None and i < len(r)) else "")

        for r in rows[1:]:
            tname = cell(r, "table")
            cname = cell(r, "column")
            if not tname or not cname:
                continue
            if tname not in tables:
                tables[tname] = {"name": tname, "columns": [], "_seen": set()}
                order.append(tname)
            b = tables[tname]
            if cname.lower() in b["_seen"]:
                continue
            b["_seen"].add(cname.lower())
            lraw = cell(r, "length")
            length = int(lraw) if lraw.isdigit() else (lraw or None)
            b["columns"].append({
                "name": cname,
                "dataType": (cell(r, "datatype") or "").lower(),
                "length": length,
                "businessTerm": cell(r, "businessterm"),
                "description": cell(r, "description"),
                "sample": cell(r, "sample"),
            })

    if not recognised_any:
        return None
    out = [{"name": t["name"], "columns": t["columns"]} for t in (tables[k] for k in order) if t["columns"]]
    return out or None


@app.route("/api/ai/extract-source", methods=["POST"])
def extract_source():
    """Read an uploaded file and use Claude to infer the SOURCE tables & columns.

    Accepts multipart/form-data with a 'file' field. Returns the standard source
    shape { tables:[{name, columns:[{name,dataType,length,businessTerm,description,sample}]}] }
    so File System sources plug into the same mapping-generation pipeline as live DBs.
    """
    up = request.files.get("file")
    if up is None:
        return jsonify(ok=False, error="No file uploaded. Attach a file in the 'file' field."), 400
    filename = up.filename or "upload"
    raw = up.read()
    if not raw:
        return jsonify(ok=False, error="The uploaded file is empty."), 400

    # Fast path #1: a STRUCTURED Excel data dictionary is parsed directly from cells —
    # instant, verbatim, no AI (falls through to AI if the layout isn't recognisable).
    if filename.lower().endswith((".xlsx", ".xlsm", ".xls")):
        xl = _parse_xlsx_dictionary(raw)
        if xl:
            cc = sum(len(t["columns"]) for t in xl)
            return jsonify(ok=True, model="xlsx-dictionary-parser", fileName=filename,
                           tables=xl, tableCount=len(xl), columnCount=cc)

    chunks, err = _extract_file_chunks(filename, raw)
    if err:
        return jsonify(ok=False, error=err), 400
    chunks = [c for c in (chunks or []) if c and c.strip()]
    if not chunks:
        return jsonify(ok=False, error="No readable text could be extracted from the file."), 400
    full_text = "\n".join(chunks)

    # Fast path #2: SQL scripts with CREATE TABLE statements are parsed deterministically
    # so EVERY table is captured (the LLM tends to summarise large DDL). If the script
    # has no parseable CREATE TABLE, fall through to the AI path below.
    if filename.lower().endswith(".sql") or "create table" in full_text.lower():
        ddl_tables = _parse_sql_ddl(full_text)
        if ddl_tables:
            col_count = sum(len(t["columns"]) for t in ddl_tables)
            return jsonify(ok=True, model="sql-ddl-parser", fileName=filename,
                           tables=ddl_tables, tableCount=len(ddl_tables), columnCount=col_count)

    if anthropic is None:
        return jsonify(ok=False, error="The 'anthropic' SDK is not installed on the server. Run: pip install anthropic"), 400

    # Cap the number of chunks (model calls) so a monster file can't run unbounded.
    truncated_chunks = len(chunks) > EXTRACT_MAX_CHUNKS
    if truncated_chunks:
        chunks = chunks[:EXTRACT_MAX_CHUNKS]

    # LOOP: extract from each chunk in its own model call, then MERGE tables by name so
    # neither input nor output truncation drops tables from large files.
    model = _ai_model()
    try:
        merged = {}          # lower(name) -> {"name":..., "columns":[...], "_cols": set()}
        order = []           # preserve first-seen table order
        for idx, chunk in enumerate(chunks):
            # one chunk failing must not abort the whole extraction — retry once, then skip
            try:
                part = _ai_extract_tables_from_text(model, filename, chunk, idx + 1, len(chunks))
            except Exception:  # noqa: BLE001
                try:
                    part = _ai_extract_tables_from_text(model, filename, chunk, idx + 1, len(chunks))
                except Exception:  # noqa: BLE001
                    part = []
            for t in part:
                key = (t.get("name") or "").strip().lower()
                if not key:
                    continue
                if key not in merged:
                    merged[key] = {"name": t["name"], "columns": [], "_cols": set()}
                    order.append(key)
                bucket = merged[key]
                for c in t.get("columns", []):
                    cn = (c.get("name") or "").strip().lower()
                    if not cn or cn in bucket["_cols"]:
                        continue   # dedup columns across chunks (e.g. repeated header)
                    bucket["_cols"].add(cn)
                    bucket["columns"].append(c)

        tables = []
        col_count = 0
        for key in order:
            b = merged[key]
            if b["columns"]:
                tables.append({"name": b["name"], "columns": b["columns"]})
                col_count += len(b["columns"])

        if not tables:
            return jsonify(ok=False, error="The AI could not identify any source tables/columns in this file."), 400

        return jsonify(ok=True, model=model, fileName=filename,
                       tables=tables, tableCount=len(tables), columnCount=col_count,
                       chunks=len(chunks), truncatedChunks=truncated_chunks)
    except Exception as exc:  # noqa: BLE001
        import traceback
        traceback.print_exc()
        msg = str(exc) or (exc.__class__.__name__ + " (see server log)")
        return jsonify(ok=False, error=msg), 400


@app.route("/api/ai/extract-source-stream", methods=["POST"])
def extract_source_stream():
    """Same as /api/ai/extract-source but STREAMS newline-delimited JSON progress
    events so the UI can show a progress bar. Events:
      {"type":"start","chunks":N,"fileName":..,"unit":"chunks|sheet-slices|pages"}
      {"type":"progress","done":i,"total":N,"tables":T,"columns":C,"label":"..."}
      {"type":"done","ok":true,"tables":[...],"tableCount":T,"columnCount":C,...}
      {"type":"error","error":".."}
    """
    up = request.files.get("file")
    if up is None:
        return jsonify(ok=False, error="No file uploaded."), 400
    filename = up.filename or "upload"
    raw = up.read()

    def gen():
        import json as _json
        def ev(obj):
            return _json.dumps(obj) + "\n"
        if not raw:
            yield ev({"type": "error", "error": "The uploaded file is empty."}); return

        # Fast path #1: structured Excel dictionary — parsed directly from cells, instant.
        if filename.lower().endswith((".xlsx", ".xlsm", ".xls")):
            xl = _parse_xlsx_dictionary(raw)
            if xl:
                cc = sum(len(t["columns"]) for t in xl)
                yield ev({"type": "start", "chunks": 1, "fileName": filename, "unit": "workbook"})
                yield ev({"type": "progress", "done": 1, "total": 1, "tables": len(xl),
                          "columns": cc, "label": "Parsed Excel dictionary (direct, no AI)"})
                yield ev({"type": "done", "ok": True, "model": "xlsx-dictionary-parser",
                          "fileName": filename, "tables": xl, "tableCount": len(xl),
                          "columnCount": cc, "chunks": 1})
                return

        chunks, err = _extract_file_chunks(filename, raw)
        if err:
            yield ev({"type": "error", "error": err}); return
        chunks = [c for c in (chunks or []) if c and c.strip()]
        if not chunks:
            yield ev({"type": "error", "error": "No readable text could be extracted from the file."}); return
        full_text = "\n".join(chunks)

        # SQL fast-path (deterministic, instant) — report as a single step.
        if filename.lower().endswith(".sql") or "create table" in full_text.lower():
            ddl = _parse_sql_ddl(full_text)
            if ddl:
                cc = sum(len(t["columns"]) for t in ddl)
                yield ev({"type": "start", "chunks": 1, "fileName": filename, "unit": "script"})
                yield ev({"type": "progress", "done": 1, "total": 1, "tables": len(ddl),
                          "columns": cc, "label": "Parsed SQL DDL"})
                yield ev({"type": "done", "ok": True, "model": "sql-ddl-parser", "fileName": filename,
                          "tables": ddl, "tableCount": len(ddl), "columnCount": cc, "chunks": 1})
                return

        if anthropic is None:
            yield ev({"type": "error", "error": "The 'anthropic' SDK is not installed on the server."}); return

        truncated = len(chunks) > EXTRACT_MAX_CHUNKS
        if truncated:
            chunks = chunks[:EXTRACT_MAX_CHUNKS]

        # infer a friendly unit label from the first chunk's header
        unit = "parts"
        head = (chunks[0].split("\n", 1)[0] if chunks else "").lower()
        if "columns " in head: unit = "column-slices"
        elif "rows " in head: unit = "row-slices"
        elif filename.lower().endswith(".pdf"): unit = "page-groups"

        yield ev({"type": "start", "chunks": len(chunks), "fileName": filename, "unit": unit})

        model = _ai_model()
        merged, order = {}, []
        try:
            failed = 0
            for idx, chunk in enumerate(chunks):
                label = chunk.split("\n", 1)[0][:80]
                # A single chunk failing (transient gateway error, etc.) must NOT abort
                # the whole extraction — retry once, then skip and keep going.
                part = []
                try:
                    part = _ai_extract_tables_from_text(model, filename, chunk, idx + 1, len(chunks))
                except Exception:  # noqa: BLE001
                    try:
                        part = _ai_extract_tables_from_text(model, filename, chunk, idx + 1, len(chunks))
                    except Exception:  # noqa: BLE001
                        failed += 1
                        part = []
                for t in part:
                    key = (t.get("name") or "").strip().lower()
                    if not key:
                        continue
                    if key not in merged:
                        merged[key] = {"name": t["name"], "columns": [], "_cols": set()}
                        order.append(key)
                    b = merged[key]
                    for c in t.get("columns", []):
                        cn = (c.get("name") or "").strip().lower()
                        if not cn or cn in b["_cols"]:
                            continue
                        b["_cols"].add(cn); b["columns"].append(c)
                tcount = sum(1 for k in order if merged[k]["columns"])
                ccount = sum(len(merged[k]["columns"]) for k in order)
                yield ev({"type": "progress", "done": idx + 1, "total": len(chunks),
                          "tables": tcount, "columns": ccount, "label": label})

            tables, col_count = [], 0
            for key in order:
                b = merged[key]
                if b["columns"]:
                    tables.append({"name": b["name"], "columns": b["columns"]})
                    col_count += len(b["columns"])
            if not tables:
                yield ev({"type": "error", "error": "The AI could not identify any tables/columns in this file."}); return
            yield ev({"type": "done", "ok": True, "model": model, "fileName": filename,
                      "tables": tables, "tableCount": len(tables), "columnCount": col_count,
                      "chunks": len(chunks), "truncatedChunks": truncated})
        except Exception as exc:  # noqa: BLE001
            import traceback; traceback.print_exc()
            yield ev({"type": "error", "error": (str(exc) or exc.__class__.__name__)})

    return app.response_class(gen(), mimetype="application/x-ndjson")


def _ai_extract_tables_from_text(model, filename, text, part_no, part_total):
    """One model call: infer source tables/columns from a single text chunk.
    Returns a normalised list of {name, columns:[...]} (empty on failure)."""
    if len(text) > EXTRACT_TEXT_BUDGET:
        text = text[:EXTRACT_TEXT_BUDGET] + "\n... (truncated)"

    system = (
        "You are a data-migration analyst. You are given part of the raw contents of a file "
        "that describes a legacy SOURCE system (a data dictionary, DDL, spec document, or a "
        "spreadsheet of ACTUAL DATA rows).\n\n"
        "Infer the source TABLES and their COLUMNS from THIS PART:\n"
        "- If it is a data dictionary or DDL, read the declared table names, column names, "
        "data types, lengths, and any descriptions or business terms.\n"
        "- If it is raw data (rows of records), treat the column headers as column names, "
        "infer each column's dataType from its values, and put one representative value in "
        "'sample'. Derive the table name from the sheet name or file name.\n"
        "- Group columns under the correct table.\n"
        "- Never invent columns that are not supported by the text.\n"
        "- EXHAUSTIVE: return EVERY table and EVERY column present in THIS PART. Do NOT "
        "summarise, sample, abbreviate, deduplicate, or omit repetitive tables.\n\n"
        "Respond with ONLY a JSON object of the form "
        '{"tables": [ {"name": "...", "columns": [ {"name": "...", "dataType": "...", '
        '"length": null, "businessTerm": "", "description": "", "sample": ""} ] } ] }. '
        "length is an integer or null. No prose, no markdown fences."
    )
    user = ("SOURCE FILE: " + filename + " (part " + str(part_no) + " of " + str(part_total) +
            ")\n\nFILE CONTENTS:\n" + text)

    client = _anthropic_client()
    base_kwargs = dict(model=model, max_tokens=16000, system=system,
                       messages=[{"role": "user", "content": user}])

    def run(extra):
        with client.messages.stream(**base_kwargs, **extra) as stream:
            return stream.get_final_message()

    attempts = [
        {"output_config": {"effort": "medium", "format": {"type": "json_schema", "schema": SOURCE_EXTRACT_SCHEMA}}},
        {"output_config": {"format": {"type": "json_schema", "schema": SOURCE_EXTRACT_SCHEMA}}},
        {"output_config": {"effort": "medium"}},
        {},
    ]
    resp, last_err = None, None
    for cfg in attempts:
        try:
            resp = run(cfg); break
        except Exception as e:  # noqa: BLE001
            last_err = e
    if resp is None:
        raise last_err
    if getattr(resp, "stop_reason", None) == "refusal":
        return []
    out = next((b.text for b in resp.content if getattr(b, "type", None) == "text"), "")
    data = _parse_mapping_json(out)
    raw_tables = data.get("tables") if isinstance(data, dict) else data
    if not isinstance(raw_tables, list):
        return []
    result = []
    for t in raw_tables:
        if not isinstance(t, dict) or not t.get("name"):
            continue
        cols = []
        for c in (t.get("columns") or []):
            if not isinstance(c, dict) or not c.get("name"):
                continue
            cols.append({
                "name": str(c.get("name", "")),
                "dataType": str(c.get("dataType", "") or ""),
                "length": c.get("length"),
                "businessTerm": str(c.get("businessTerm", "") or ""),
                "description": str(c.get("description", "") or ""),
                "sample": c.get("sample", ""),
            })
        if cols:
            result.append({"name": str(t["name"]), "columns": cols})
    return result


SOURCE_EXTRACT_SCHEMA = {
    "type": "object",
    "properties": {
        "tables": {
            "type": "array",
            "items": {
                "type": "object",
                "properties": {
                    "name": {"type": "string"},
                    "columns": {
                        "type": "array",
                        "items": {
                            "type": "object",
                            "properties": {
                                "name": {"type": "string"},
                                "dataType": {"type": "string"},
                                "length": {"type": ["integer", "null"]},
                                "businessTerm": {"type": "string"},
                                "description": {"type": "string"},
                                "sample": {"type": "string"},
                            },
                            "required": ["name", "dataType"],
                            "additionalProperties": False,
                        },
                    },
                },
                "required": ["name", "columns"],
                "additionalProperties": False,
            },
        },
    },
    "required": ["tables"],
    "additionalProperties": False,
}


if __name__ == "__main__":
    port = int(os.environ.get("PORT", 8000))
    app.run(host="127.0.0.1", port=port, debug=True)
